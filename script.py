
"""
At the command line, only need to run once to install the package via pip:

$ pip install pillow google-generativeai python-dotenv python-docx 
"""
# Import necessary libraries
import os
import google.generativeai as genai
from PIL import Image
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from IPython.display import Markdown

# Load API key from environment variables
user_secrets = load_dotenv()
api_key = os.getenv('GEMINI_API_KEY')

# Configure the generative AI model with the API key
genai.configure(api_key=api_key)

# Configuration settings for text generation
generation_config = {
  "temperature": 0.9,
  "top_p": 1,
  "top_k": 32,
  "max_output_tokens": 7096,
}

# Safety settings to control content generation
safety_settings = [
  {
    "category": "HARM_CATEGORY_HARASSMENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_HATE_SPEECH",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  },
  {
    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
  }
]

# Initialize the generative model
model = genai.GenerativeModel(
    model_name="gemini-pro-vision",
    generation_config=generation_config,
    safety_settings=safety_settings
)

# Instruction for image comparison
instruction = """
Compare two images, the base image is a Figma design, the other is the final developed product.
To check the UI, compare things like icons, logos, text box, colors, and so on.
If there are no differences, just mention 'Both images are the same' and provide a precise explanation of the difference. Only give me differences.
"""

def get_directories():
    """Create directories if they don't exist and collect file paths.

    This function creates 'base_dir/' and 'target_dir/' directories if they don't exist,
    and then retrieves the file paths within these directories.

    Returns:
        tuple: A tuple containing two lists - BASE_FILES and TARGET_FILES.
        BASE_FILES (list): Contains file paths within 'base_dir/'.
        TARGET_FILES (list): Contains file paths within 'target_dir/'.
    """

    BASE_DIR = 'base_dir/'
    TARGET_DIR = 'target_dir/'
    
    directories = [BASE_DIR, TARGET_DIR]
    
    for directory in directories:
        os.makedirs(directory, exist_ok=True)
        
    BASE_FILES = []
    TARGET_FILES = []
    
    for directory, files_list in zip(directories, [BASE_FILES, TARGET_FILES]):
        files_list.extend([os.path.join(root, filename) for root, dirs, files in os.walk(directory) for filename in files])
    
    return BASE_FILES, TARGET_FILES

# Get directories and file paths for base and target images
base_files, target_files = get_directories()

# Create a Word document to store comparison results
doc = Document()

doc.add_heading('Image Comparison Report', level=1)

# Iterate through base and target images to compare
for base_image, target_image in zip(base_files, target_files):
    print(f"base image is:{base_image}, Target image is {target_image}")
    # Open base and target images
    base_img, target_img = Image.open(base_image), Image.open(target_image)
    prompt = [instruction, base_img, target_img]
    
    # Generate content based on comparison
    response = model.generate_content(prompt)
    
    # Add comparison results to the Word document
    doc.add_paragraph(f"Base image: {base_image}, Target image: {target_image}")
    text_content = response.text
    doc.add_paragraph(text_content)
    doc.add_paragraph("--------------------------------------")

# Save the comparison report in a Word document
doc.save('Comparison_Report.docx')
